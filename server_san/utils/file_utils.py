import fitz  # PyMuPDF
from docx import Document
import io

def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats using file path"""
    file_extension = file_path.split('.')[-1].lower()
    text = ''
    
    if file_extension == 'pdf':
        # Extract text from PDF using PyMuPDF
        with open(file_path, 'rb') as file:
            doc = fitz.open(stream=file.read(), filetype='pdf')
            for page in doc:
                text += page.get_text("text")
    elif file_extension == 'docx':
        # Extract text from DOCX
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    elif file_extension == 'txt':
        # Extract text from TXT
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    else:
        raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
    return text

def extract_text_from_bytes(file_content: bytes, file_extension: str) -> str:
    """Extract text from various file formats using file bytes"""
    text = ''
    
    if file_extension == 'pdf':
        # Extract text from PDF using PyMuPDF
        doc = fitz.open(stream=file_content, filetype='pdf')
        for page in doc:
            text += page.get_text("text")
    elif file_extension == 'docx':
        # Extract text from DOCX
        doc = Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    elif file_extension == 'txt':
        # Extract text from TXT
        text = file_content.decode('utf-8')
    else:
        raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
    return text